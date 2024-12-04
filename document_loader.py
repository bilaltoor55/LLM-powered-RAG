import os
from typing import List
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    DirectoryLoader,
    PyPDFLoader,
    TextLoader,
)
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from docx import Document as DocxDocument

# Constants
TEXT_SPLITTER = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)

def load_documents_into_database(model_name: str, documents_path: str) -> Chroma:
    """
    Loads documents from the specified path (directory or file) into the Chroma database
    after splitting the text into chunks.

    Args:
        model_name (str): The name of the embedding model.
        documents_path (str): Path to the directory or file containing documents.

    Returns:
        Chroma: The Chroma database with loaded documents.
    """
    # Resolve absolute path
    documents_path = os.path.abspath(documents_path)
    
    if not os.path.exists(documents_path):
        raise FileNotFoundError(f"The specified path does not exist: {documents_path}")

    print(f"Loading documents from: {documents_path}")
    raw_documents = load_documents(documents_path)
    documents = TEXT_SPLITTER.split_documents(raw_documents)

    print("Creating embeddings and loading documents into Chroma")
    db = Chroma.from_documents(
        documents,
        OllamaEmbeddings(model=model_name),
    )
    print("Documents successfully loaded into the Chroma database.")
    return db


def load_documents(path: str) -> List[Document]:
    """
    Loads documents from the specified path, which can be either a directory or a single file.

    Args:
        path (str): The path to the directory or file containing documents to load.

    Returns:
        List[Document]: A list of loaded documents.

    Raises:
        FileNotFoundError: If the specified path does not exist.
        ValueError: If the specified path is invalid or unsupported file types are provided.
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"The specified path does not exist: {path}")

    docs = []

    # If the path is a directory, use DirectoryLoader
    if os.path.isdir(path):
        loaders = {
            ".pdf": DirectoryLoader(
                path,
                glob="**/*.pdf",
                loader_cls=PyPDFLoader,
                show_progress=True,
                use_multithreading=True,
            ),
            ".md": DirectoryLoader(
                path,
                glob="**/*.md",
                loader_cls=TextLoader,
                show_progress=True,
            ),
            ".txt": DirectoryLoader(
                path,
                glob="**/*.txt",
                loader_cls=TextLoader,
                show_progress=True,
            ),
            ".docx": DirectoryLoader(
                path,
                glob="**/*.docx",
                loader_cls=TextLoader,
                show_progress=True,
            ),
        }

        for file_type, loader in loaders.items():
            print(f"Loading {file_type} files from directory: {path}")
            docs.extend(loader.load())

    # If the path is a single file, handle file types directly
    elif os.path.isfile(path):
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            print(f"Loading single PDF file: {path}")
            loader = PyPDFLoader(path)
            docs.extend(loader.load())
        elif ext == ".md":
            print(f"Loading single Markdown file: {path}")
            loader = TextLoader(path)
            docs.extend(loader.load())
        elif ext == ".txt":
            print(f"Loading single TXT file: {path}")
            with open(path, "r") as file:
                content = file.read()
            # Adding metadata
            docs.append(Document(page_content=content, metadata={"source": path, "page": 1}))
        elif ext == ".docx":
            print(f"Loading single DOCX file: {path}")
            doc = DocxDocument(path)
            content = ""
            for para in doc.paragraphs:
                content += para.text + "\n"
            # Adding metadata
            docs.append(Document(page_content=content, metadata={"source": path, "page": 1}))
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    else:
        raise ValueError(f"Invalid path: {path}")

    if not docs:
        raise ValueError(f"No valid documents found at the specified path: {path}")

    return docs


if __name__ == "__main__":
    # Example usage for testing
    model_name = "llama-2"  # Replace with your embedding model name
    documents_path = "/root/LLM-powered-RAG/Research"
    # Replace with the correct path to your files "/path/to/documents"  

    try:
        db = load_documents_into_database(model_name, documents_path)
        print("Chroma database initialized successfully.")
    except Exception as e:
        print(f"Error: {e}")
